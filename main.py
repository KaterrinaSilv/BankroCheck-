import random
import re
from docx import Document
from faker import Faker

fake = Faker('ru_RU')


def generate_full_name():
    gender = random.choice(['male', 'female'])

    if gender == 'male':
        first_name = fake.first_name_male()
        last_name = fake.last_name_male()
        middle_name = fake.middle_name_male()
    else:
        first_name = fake.first_name_female()
        last_name = fake.last_name_female()
        middle_name = fake.middle_name_female()

    return {
        "имя": first_name,
        "фамилия": last_name,
        "отчество": middle_name
    }


def replace_tags(text, full_name_data):
    tags = ["имя", "фамилия", "отчество"]

    for tag in tags:
        pattern = r"\{\{\s*" + tag + r"\s*\}\}"
        if re.search(pattern, text):
            text = re.sub(pattern, full_name_data[tag], text)

    return text


def process_tables(doc, full_name_data):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = replace_tags(cell.text, full_name_data)
                # print(cell.text)


def fill_document(input_file, output_file):
    doc = Document(input_file)
    full_name_data = generate_full_name()

    for paragraph in doc.paragraphs:
        paragraph.text = replace_tags(paragraph.text, full_name_data)

    process_tables(doc, full_name_data)

    doc.save(output_file)


input_file = "template.docx"  # Входной файл с шаблоном
output_file = "filled_document.docx"  # Выходной заполненный файл

fill_document(input_file, output_file)
